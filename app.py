import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO

# CONFIGURACIÓN DE PÁGINA
st.set_page_config(page_title="Valorización Día Operativo - Oficial", page_icon="🛡️", layout="wide")

st.title("🛡️ Sistema de Valorización de Costos Operativos (SRT)")
st.info("Cómputo basado en Orden de Valorización de Medios y Personal - Datos Oficiales")

# --- BASE DE DATOS TÉCNICA ---
db_vehiculos = {
    "Ford Ranger (Gasoil)": {"cons_100": 11.0, "mant_km": 150},
    "Hummvee / Hummer (Gasoil)": {"cons_100": 30.0, "mant_km": 450},
    "Jeep MB 230G (Nafta)": {"cons_100": 25.0, "mant_km": 300},
    "Unimog (Gasoil)": {"cons_100": 20.0, "mant_km": 400}
}

db_viaticos = {
    "JEMGE": 114644.60, "Oficial Superior": 94581.80,
    "Oficial Jefe": 74518.99, "Oficial Subalterno": 63054.53,
    "Suboficial Superior": 57322.30, "Suboficial Subalterno": 51590.07,
    "Cadete": 37259.50, "Soldado Voluntario / Aspirante": 34393.38
}

db_drones_usd = {
    "Mavic II Enterprise": 3150.0, "Mavic IV": 6950.0,
    "Phantom 4 Pro": 1260.0, "Phantom 4 RTK": 7890.0, "Matrice 200": 5700.0
}

# --- BARRA LATERAL ---
with st.sidebar:
    st.header("📊 Mercado y Divisas")
    tc_bna = st.number_input("Tipo de Cambio BNA (Vendedor)", min_value=1.0, value=950.0)
    
    st.header("⛽ Combustible (Manual)")
    p_nafta = st.number_input("Precio Litro Nafta ($)", min_value=0.0, value=1114.0)
    p_gasoil = st.number_input("Precio Litro Gasoil ($)", min_value=0.0, value=1414.0)

    st.divider()
    st.header("📡 Conectividad (Starlink)")
    posee_antena = st.checkbox("¿Ya posee antena Starlink?", value=False)
    costo_antena = 300000.0 # Según tu planilla
    abono_mensual = 87500.0 # Según tu planilla
    
    st.divider()
    st.header("⚙️ Configuración Misión")
    vehiculo_sel = st.selectbox("Vehículo 4x4", list(db_vehiculos.keys()))
    km_despliegue = st.number_input("Distancia Despliegue/Repliegue (km)", value=100)
    
    st.subheader("👥 Personal")
    p1 = st.selectbox("Conductor", list(db_viaticos.keys()), index=5)
    p2 = st.selectbox("Operador 1", list(db_viaticos.keys()), index=3)
    p3 = st.selectbox("Operador 2", list(db_viaticos.keys()), index=4)
    p4 = st.selectbox("Auxiliar", list(db_viaticos.keys()), index=7)
    
    dron_sel = st.selectbox("Modelo de Dron (USD)", list(db_drones_usd.keys()))
    horas_gen = st.slider("Horas Generador (Dogo 3500)", 0, 24, 8)

# --- CÁLCULOS ---
# 1. Despliegue
v_data = db_vehiculos[vehiculo_sel]
p_uso = p_gasoil if "Gasoil" in vehiculo_sel else p_nafta
costo_comb_dr = (km_despliegue / 100) * v_data["cons_100"] * p_uso
costo_mant_dr = km_despliegue * v_data["mant_km"]
total_dr_ars = costo_comb_dr + costo_mant_dr

# 2. Operativo
viaticos_ars = db_viaticos[p1] + db_viaticos[p2] + db_viaticos[p3] + db_viaticos[p4]
comb_gen_ars = horas_gen * 1.5 * p_gasoil

# Sección Starlink
costo_dia_starlink = abono_mensual / 30
inversion_antena = 0.0 if posee_antena else costo_antena

# 3. Equipo
valor_usd = db_drones_usd[dron_sel]
valor_ars_bna = valor_usd * tc_bna
amortizacion_ars = valor_ars_bna * 0.001

total_op_ars = viaticos_ars + comb_gen_ars + costo_dia_starlink + amortizacion_ars
total_mision = total_dr_ars + total_op_ars + inversion_antena

# --- INTERFAZ ---
st.markdown(f"### 💹 Cotización: **$ {tc_bna:,.2f}** (BNA Vendedor)")

# Resumen General
col_a, col_b, col_c = st.columns(3)
col_a.metric("DESPLIEGUE/REPLIEGUE", f"ARS {total_dr_ars:,.2f}")
col_b.metric("COSTO DÍA OPERATIVO", f"ARS {total_op_ars:,.2f}")
if not posee_antena:
    col_c.metric("INVERSIÓN ÚNICA VEZ", f"ARS {inversion_antena:,.2f}")
else:
    col_c.metric("EQUIPO SRT (USD)", f"u$s {valor_usd:,.2f}")

st.divider()

# TABLA DE DETALLE
resumen_data = [
    ["Combustible Movilidad", "ARS", f"$ {costo_comb_dr:,.2f}", costo_comb_dr],
    ["Mantenimiento Programado", "ARS", f"$ {costo_mant_dr:,.2f}", costo_mant_dr],
    ["Viáticos Personal (100%)", "ARS", f"$ {viaticos_ars:,.2f}", viaticos_ars],
    ["Energía (Generador)", "ARS", f"$ {comb_gen_ars:,.2f}", comb_gen_ars],
    ["Servicio Starlink (Día Itinerante)", "ARS", f"$ {costo_dia_starlink:,.2f}", costo_dia_starlink],
    ["Amortización Equipo SRT", "USD", f"u$s {valor_usd:,.2f}", amortizacion_ars]
]

if not posee_antena:
    resumen_data.append(["Adquisición Antena Starlink (Única vez)", "ARS", f"$ {costo_antena:,.2f}", costo_antena])

df_resumen = pd.DataFrame(resumen_data, columns=["Rubro", "Origen", "Costo Detalle", "Subtotal ARS"])
st.subheader("📋 Detalle de Valorización (Consolidado)")
st.table(df_resumen.drop(columns=["Subtotal ARS"]))

st.header(f"TOTAL VALORIZADO: $ {total_mision:,.2f}")

# --- FUNCIÓN GENERAR WORD ---
def generate_docx(df, total_final):
    doc = Document()
    doc.add_heading('PLANILLA DE VALORIZACIÓN DE COSTOS - SRT', 0)
    
    p = doc.add_paragraph()
    p.add_run(f"Vehículo: {vehiculo_sel}\n").bold = True
    p.add_run(f"Tipo de Cambio BNA: {tc_bna}\n")
    p.add_run(f"Dron: {dron_sel}\n")

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Rubro'
    hdr_cells[1].text = 'Origen'
    hdr_cells[2].text = 'Costo (ARS)'

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Rubro'])
        row_cells[1].text = str(row['Origen'])
        row_cells[2].text = f"$ {row['Subtotal ARS']:,.2f}"

    doc.add_paragraph(f"\nTOTAL FINAL VALORIZADO: $ {total_final:,.2f}").bold = True
    doc.add_paragraph("\nNota: La adquisición de antena Starlink se considera un gasto de capital por única vez.")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# BOTÓN DESCARGA
st.divider()
doc_buffer = generate_docx(df_resumen, total_mision)
st.download_button(
    label="📥 Descargar Planilla en WORD (Editable)",
    data=doc_buffer,
    file_name="valorizacion_mision_srt.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
